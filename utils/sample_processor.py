"""
Sample-Based Grant Generation

This script processes your grant samples and uses them to generate new grants
that match the style, structure, and approach of successful proposals.
Supports text files, PDFs, and Word documents.
"""

import os
import glob
import io
import re
import time
from typing import List, Dict, Any, Optional, Tuple, Union
from pathlib import Path
import logging
from datetime import datetime

# Setup logging
logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Import document processing libraries
try:
    import pypdf
    PDF_SUPPORT = True
except ImportError:
    logger.warning("pypdf not found. PDF support will be disabled.")
    PDF_SUPPORT = False

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    logger.warning(
        "python-docx not found. Word document support will be disabled.")
    DOCX_SUPPORT = False


class SampleGrantProcessor:
    """Process grant samples to extract style and structure for generation."""

    def __init__(self, samples_dir: str = "data/samples"):
        self.samples_dir = Path(samples_dir)
        self.grant_types = ["research", "nonprofit",
                            "technology", "healthcare", "education"]
        self.samples = {}
        self.supported_extensions = [".txt"]

        # Add supported extensions based on available libraries
        if PDF_SUPPORT:
            self.supported_extensions.append(".pdf")
        if DOCX_SUPPORT:
            self.supported_extensions.extend([".docx", ".doc"])

        logger.info(
            f"Initialized SampleGrantProcessor with support for: {', '.join(self.supported_extensions)}")

    def _extract_text_from_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text and metadata from a PDF file."""
        if not PDF_SUPPORT:
            raise ImportError(
                "PDF support is not available. Please install pypdf.")

        metadata = {}
        text = ""

        try:
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = pypdf.PdfReader(pdf_file)

                # Extract metadata if available
                if pdf_reader.metadata:
                    meta = pdf_reader.metadata
                    metadata = {
                        "author": meta.author if hasattr(meta, 'author') else None,
                        "creation_date": meta.creation_date if hasattr(meta, 'creation_date') else None,
                        "creator": meta.creator if hasattr(meta, 'creator') else None,
                        "producer": meta.producer if hasattr(meta, 'producer') else None,
                        "title": meta.title if hasattr(meta, 'title') else None,
                        "page_count": len(pdf_reader.pages)
                    }

                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                    else:
                        logger.warning(
                            f"Page {page_num+1} in {file_path.name} contains no extractable text")

                logger.info(
                    f"Extracted {len(text.split())} words from PDF: {file_path.name}")

        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise

        return text, metadata

    def _extract_text_from_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text and metadata from a Word document."""
        if not DOCX_SUPPORT:
            raise ImportError(
                "Word document support is not available. Please install python-docx.")

        metadata = {}
        text = ""

        try:
            doc = docx.Document(file_path)

            # Extract metadata if available
            core_properties = doc.core_properties
            metadata = {
                "author": core_properties.author,
                "created": core_properties.created,
                "modified": core_properties.modified,
                "title": core_properties.title,
                "paragraph_count": len(doc.paragraphs)
            }

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                text += "\n"

            logger.info(
                f"Extracted {len(text.split())} words from Word document: {file_path.name}")

        except Exception as e:
            logger.error(
                f"Error extracting text from Word document {file_path}: {e}")
            raise

        return text, metadata

    def _extract_text_from_file(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text and metadata from a file based on its extension."""
        extension = file_path.suffix.lower()

        # Check if file extension is supported
        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file extension: {extension}")

        # Process based on file type
        if extension == ".pdf" and PDF_SUPPORT:
            return self._extract_text_from_pdf(file_path)
        elif extension in [".docx", ".doc"] and DOCX_SUPPORT:
            return self._extract_text_from_docx(file_path)
        elif extension == ".txt":
            # For text files, just read the content
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()

                # Basic metadata for text files
                metadata = {
                    "created": datetime.fromtimestamp(os.path.getctime(file_path)),
                    "modified": datetime.fromtimestamp(os.path.getmtime(file_path)),
                    "size_bytes": os.path.getsize(file_path)
                }

                return content, metadata

            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                raise
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    def _validate_and_sanitize_text(self, text: str) -> str:
        """Validate and sanitize extracted text."""
        if not text or len(text.strip()) == 0:
            logger.warning(
                "Extracted text is empty or contains only whitespace")
            return ""

        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)

        # Remove any control characters
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)

        # Ensure proper line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        return text.strip()

    def load_samples(self, show_progress: bool = False) -> Dict[str, List[Dict[str, Any]]]:
        """Load all grant samples from the samples directory."""
        samples = {}
        total_files = 0
        processed_files = 0

        # Count total files first if showing progress
        if show_progress:
            for grant_type in self.grant_types:
                type_dir = self.samples_dir / grant_type
                if type_dir.exists():
                    for ext in self.supported_extensions:
                        total_files += len(list(type_dir.glob(f"*{ext}")))

        start_time = time.time()

        for grant_type in self.grant_types:
            type_dir = self.samples_dir / grant_type
            if type_dir.exists():
                samples[grant_type] = []

                # Process all supported file types
                for ext in self.supported_extensions:
                    for file_path in type_dir.glob(f"*{ext}"):
                        try:
                            # Extract text and metadata
                            content, metadata = self._extract_text_from_file(
                                file_path)

                            # Validate and sanitize the text
                            content = self._validate_and_sanitize_text(content)
                            if not content:
                                logger.warning(
                                    f"Skipping empty file: {file_path}")
                                continue

                            # Add to samples
                            samples[grant_type].append({
                                "filename": file_path.name,
                                "content": content,
                                "word_count": len(content.split()),
                                "sections": self._extract_sections(content),
                                # Remove the dot
                                "file_type": file_path.suffix.lower()[1:],
                                "metadata": metadata
                            })

                            logger.info(
                                f"Loaded sample: {file_path.name} ({file_path.suffix.lower()[1:]})")

                            # Update progress
                            processed_files += 1
                            if show_progress and total_files > 0:
                                progress = (processed_files /
                                            total_files) * 100
                                elapsed = time.time() - start_time
                                logger.info(
                                    f"Progress: {progress:.1f}% ({processed_files}/{total_files}) - {elapsed:.1f}s elapsed")

                        except Exception as e:
                            logger.error(f"Error loading {file_path}: {e}")

        self.samples = samples

        # Log summary
        total_loaded = sum(len(samples_list)
                           for samples_list in samples.values())
        logger.info(
            f"Loaded {total_loaded} samples in {time.time() - start_time:.2f}s")

        return samples

    def _extract_sections(self, content: str) -> List[str]:
        """Extract section headings from the grant content."""
        lines = content.split('\n')
        sections = []

        for line in lines:
            line = line.strip()
            # Common grant section patterns
            if (line.isupper() or
                line.startswith('##') or
                line.startswith('TITLE:') or
                line.startswith('EXECUTIVE SUMMARY') or
                line.startswith('PROBLEM STATEMENT') or
                line.startswith('METHODOLOGY') or
                line.startswith('EXPECTED OUTCOMES') or
                line.startswith('BUDGET') or
                'SUMMARY' in line.upper() or
                'STATEMENT' in line.upper() or
                'METHODOLOGY' in line.upper() or
                    'OUTCOMES' in line.upper()):
                sections.append(line)

        return sections

    def get_best_sample(self, grant_type: str, topic: str = None) -> Optional[Dict[str, Any]]:
        """Get the best matching sample for a given grant type and topic."""
        if grant_type not in self.samples or not self.samples[grant_type]:
            # Fallback to any available sample
            for gtype, samples_list in self.samples.items():
                if samples_list:
                    logger.info(
                        f"Using {gtype} sample as fallback for {grant_type}")
                    return samples_list[0]
            return None

        # For now, return the first sample of the requested type
        # TODO: Add topic similarity matching
        return self.samples[grant_type][0]

    def create_style_prompt(self, sample: Dict[str, Any], user_topic: str, section: str) -> str:
        """Create a prompt that incorporates the sample's style and structure."""
        if not sample:
            return f"Write a {section} for a grant proposal about {user_topic}."

        # Extract relevant sections from the sample
        content = sample["content"]
        sections = sample["sections"]

        # Find the matching section in the sample
        sample_section = ""
        for sec in sections:
            if section.upper() in sec.upper():
                # Extract the content under this section
                start_idx = content.find(sec)
                if start_idx != -1:
                    # Find the next section or end of content
                    remaining_content = content[start_idx + len(sec):]
                    next_section_idx = float('inf')

                    for other_sec in sections:
                        if other_sec != sec:
                            idx = remaining_content.find(other_sec)
                            if idx != -1 and idx < next_section_idx:
                                next_section_idx = idx

                    if next_section_idx == float('inf'):
                        sample_section = remaining_content
                    else:
                        sample_section = remaining_content[:next_section_idx]

                    sample_section = sample_section.strip()
                    break

        # Create the style-aware prompt
        if sample_section:
            prompt = f"""
            Write a {section} for a grant proposal about "{user_topic}".
            
            Use this successful example as a style and structure guide:
            
            EXAMPLE {section.upper()}:
            {sample_section[:800]}...
            
            Match the writing style, tone, and structure of the example while adapting the content 
            to focus on "{user_topic}". Maintain the same level of detail and professional language.
            """
        else:
            # Fallback to general style guidance
            first_500_words = " ".join(content.split()[:500])
            prompt = f"""
            Write a {section} for a grant proposal about "{user_topic}".
            
            Use this successful grant proposal as a style guide:
            
            STYLE EXAMPLE:
            {first_500_words}...
            
            Match the writing style, tone, and level of formality while focusing on "{user_topic}".
            """

        return prompt

    def get_sample_statistics(self) -> Dict[str, Any]:
        """Get statistics about loaded samples."""
        stats = {
            "total_samples": sum(len(samples) for samples in self.samples.values()),
            "by_type": {},
            "average_length": 0,
            "by_file_type": {}
        }

        total_words = 0
        total_samples = 0
        file_type_counts = {}

        for grant_type, samples_list in self.samples.items():
            if samples_list:
                stats["by_type"][grant_type] = {
                    "count": len(samples_list),
                    "avg_words": sum(s["word_count"] for s in samples_list) / len(samples_list),
                    "sections": [],
                    "file_types": {}
                }

                # Collect all sections
                all_sections = []
                for sample in samples_list:
                    all_sections.extend(sample["sections"])

                    # Count file types
                    file_type = sample.get("file_type", "unknown")
                    if file_type not in stats["by_type"][grant_type]["file_types"]:
                        stats["by_type"][grant_type]["file_types"][file_type] = 0
                    stats["by_type"][grant_type]["file_types"][file_type] += 1

                    # Global file type stats
                    if file_type not in file_type_counts:
                        file_type_counts[file_type] = 0
                    file_type_counts[file_type] += 1

                stats["by_type"][grant_type]["sections"] = list(
                    set(all_sections))

                total_words += sum(s["word_count"] for s in samples_list)
                total_samples += len(samples_list)

        # Calculate overall file type statistics
        for file_type, count in file_type_counts.items():
            stats["by_file_type"][file_type] = {
                "count": count,
                "percentage": (count / total_samples * 100) if total_samples > 0 else 0
            }

        if total_samples > 0:
            stats["average_length"] = total_words / total_samples

        return stats

    def batch_process_samples(self, batch_size: int = 5) -> None:
        """Process samples in batches to avoid memory issues with large collections."""
        all_files = []

        # Collect all files first
        for grant_type in self.grant_types:
            type_dir = self.samples_dir / grant_type
            if type_dir.exists():
                for ext in self.supported_extensions:
                    all_files.extend(list(type_dir.glob(f"*{ext}")))

        total_files = len(all_files)
        logger.info(
            f"Found {total_files} files to process in batches of {batch_size}")

        # Process in batches
        for i in range(0, total_files, batch_size):
            batch = all_files[i:i+batch_size]
            logger.info(
                f"Processing batch {i//batch_size + 1}/{(total_files + batch_size - 1)//batch_size}")

            for file_path in batch:
                try:
                    grant_type = file_path.parent.name
                    if grant_type in self.grant_types:
                        # Extract text and metadata
                        content, metadata = self._extract_text_from_file(
                            file_path)

                        # Validate and sanitize
                        content = self._validate_and_sanitize_text(content)
                        if not content:
                            continue

                        # Make sure the grant type exists in samples
                        if grant_type not in self.samples:
                            self.samples[grant_type] = []

                        # Add to samples
                        self.samples[grant_type].append({
                            "filename": file_path.name,
                            "content": content,
                            "word_count": len(content.split()),
                            "sections": self._extract_sections(content),
                            "file_type": file_path.suffix.lower()[1:],
                            "metadata": metadata
                        })

                except Exception as e:
                    logger.error(f"Error in batch processing {file_path}: {e}")

    def add_sample_to_rag(self, rag_utils):
        """Add all samples to the RAG system for enhanced context retrieval."""
        try:
            from utils.rag_utils import store_text_content

            for grant_type, samples_list in self.samples.items():
                for sample in samples_list:
                    # Create a temporary file-like object
                    text_file = io.StringIO(sample["content"])
                    text_file.name = f"{grant_type}_{sample['filename']}"

                    # Store in RAG system
                    store_text_content(
                        text_file, f"Sample {grant_type} grant: {sample['filename']}")
                    logger.info(f"Added {sample['filename']} to RAG system")

        except Exception as e:
            logger.error(f"Error adding samples to RAG: {e}")

# Usage functions


def process_user_samples(show_progress: bool = True):
    """Process user samples and return the processor."""
    processor = SampleGrantProcessor()
    samples = processor.load_samples(show_progress=show_progress)

    print(f"📊 Loaded {sum(len(s) for s in samples.values())} grant samples")

    stats = processor.get_sample_statistics()
    print("\n📊 Sample Statistics:")
    for grant_type, info in stats["by_type"].items():
        print(
            f"  {grant_type}: {info['count']} samples, avg {info['avg_words']:.0f} words")

    # Print file type statistics
    print("\n📄 File Types:")
    for file_type, info in stats["by_file_type"].items():
        print(
            f"  {file_type}: {info['count']} files ({info['percentage']:.1f}%)")

    return processor


def generate_with_sample_style(processor, grant_type: str, topic: str, section: str):
    """Generate a proposal section using sample style."""
    sample = processor.get_best_sample(grant_type, topic)

    if sample:
        print(f"✅ Using sample: {sample['filename']} for {grant_type} grant")
        style_prompt = processor.create_style_prompt(sample, topic, section)
        return style_prompt
    else:
        print(f"❌ No samples found for {grant_type}")
        return f"Write a {section} for a grant proposal about {topic}."


if __name__ == "__main__":
    # Demo usage
    processor = process_user_samples()

    # Example: Generate a problem statement using research grant style
    prompt = generate_with_sample_style(
        processor,
        "research",
        "AI-powered healthcare diagnostics",
        "Problem Statement"
    )

    print("\n📝 Generated Style-Aware Prompt:")
    print(prompt)
