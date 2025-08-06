"""
Test Enhanced Sample Processor

This script tests the enhanced sample processor with its new capabilities
for processing PDF and Word documents in addition to text files.
"""

from utils.sample_processor import SampleGrantProcessor, PDF_SUPPORT, DOCX_SUPPORT
import os
import time
import logging
import shutil
from pathlib import Path
from typing import List, Dict, Any
import random

# Setup logging
logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Import the enhanced sample processor

# Check for required libraries
try:
    import pypdf
except ImportError:
    logger.warning("pypdf not installed. PDF tests will be skipped.")

try:
    import docx
    from docx.shared import Pt
except ImportError:
    logger.warning(
        "python-docx not installed. Word document tests will be skipped.")

# Test data directory
TEST_DATA_DIR = Path("data/test_samples")
GRANT_TYPES = ["research", "nonprofit",
               "technology", "healthcare", "education"]
SAMPLE_SECTIONS = [
    "EXECUTIVE SUMMARY",
    "PROBLEM STATEMENT",
    "METHODOLOGY",
    "EXPECTED OUTCOMES",
    "BUDGET",
    "TIMELINE"
]


def create_test_files() -> Dict[str, List[Path]]:
    """Create test files of different types if they don't exist."""
    logger.info("Creating test files...")

    # Create test directory structure
    if TEST_DATA_DIR.exists():
        logger.info(
            f"Test directory {TEST_DATA_DIR} already exists. Clearing it...")
        shutil.rmtree(TEST_DATA_DIR)

    # Create directories for each grant type
    created_files = {}
    for grant_type in GRANT_TYPES:
        type_dir = TEST_DATA_DIR / grant_type
        type_dir.mkdir(parents=True, exist_ok=True)
        created_files[grant_type] = []

        # Create text files
        logger.info(f"Creating text files for {grant_type}...")
        for i in range(2):
            file_path = type_dir / f"sample_{grant_type}_{i+1}.txt"
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"# {grant_type.upper()} GRANT PROPOSAL {i+1}\n\n")

                # Add random sections
                for section in SAMPLE_SECTIONS:
                    f.write(f"\n{section}\n")
                    f.write(
                        f"This is sample content for the {section.lower()} section of a {grant_type} grant proposal.\n")
                    f.write(
                        "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.\n")
                    f.write(
                        "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.\n\n")

            created_files[grant_type].append(file_path)

        # Create PDF files if supported
        if PDF_SUPPORT:
            logger.info(f"Creating PDF files for {grant_type}...")
            try:
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter

                for i in range(2):
                    file_path = type_dir / f"sample_{grant_type}_{i+1}.pdf"
                    c = canvas.Canvas(str(file_path), pagesize=letter)
                    c.setFont("Helvetica", 12)

                    # Add title
                    c.drawString(
                        72, 750, f"{grant_type.upper()} GRANT PROPOSAL {i+1}")

                    # Add sections
                    y_position = 700
                    for section in SAMPLE_SECTIONS:
                        c.drawString(72, y_position, section)
                        y_position -= 20
                        c.drawString(
                            72, y_position, f"This is sample content for the {section.lower()} section of a {grant_type} grant proposal.")
                        y_position -= 15
                        c.drawString(
                            72, y_position, "Lorem ipsum dolor sit amet, consectetur adipiscing elit.")
                        y_position -= 15
                        c.drawString(
                            72, y_position, "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris.")
                        y_position -= 30

                    c.save()
                    created_files[grant_type].append(file_path)

            except Exception as e:
                logger.error(f"Error creating PDF files: {e}")

        # Create Word documents if supported
        if DOCX_SUPPORT:
            logger.info(f"Creating Word documents for {grant_type}...")
            try:
                for i in range(2):
                    file_path = type_dir / f"sample_{grant_type}_{i+1}.docx"
                    doc = docx.Document()

                    # Add title
                    title = doc.add_heading(
                        f"{grant_type.upper()} GRANT PROPOSAL {i+1}", level=1)

                    # Add sections
                    for section in SAMPLE_SECTIONS:
                        doc.add_heading(section, level=2)
                        doc.add_paragraph(
                            f"This is sample content for the {section.lower()} section of a {grant_type} grant proposal.")
                        p = doc.add_paragraph(
                            "Lorem ipsum dolor sit amet, consectetur adipiscing elit. ")
                        p.add_run(
                            "Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.").bold = True
                        doc.add_paragraph(
                            "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.")

                    doc.save(file_path)
                    created_files[grant_type].append(file_path)

            except Exception as e:
                logger.error(f"Error creating Word documents: {e}")

    logger.info(f"Created test files in {TEST_DATA_DIR}")
    return created_files


def test_file_loading():
    """Test loading files of different types."""
    logger.info("\n===== TESTING FILE LOADING =====")

    # Initialize processor with test directory
    processor = SampleGrantProcessor(samples_dir=str(TEST_DATA_DIR))

    # Time the loading process
    start_time = time.time()
    samples = processor.load_samples(show_progress=True)
    elapsed = time.time() - start_time

    # Print summary
    total_samples = sum(len(s) for s in samples.values())
    logger.info(f"Loaded {total_samples} samples in {elapsed:.2f} seconds")

    # Print breakdown by file type
    stats = processor.get_sample_statistics()
    logger.info("\nSample Statistics:")
    for grant_type, info in stats["by_type"].items():
        logger.info(
            f"  {grant_type}: {info['count']} samples, avg {info['avg_words']:.0f} words")
        for file_type, count in info["file_types"].items():
            logger.info(f"    - {file_type}: {count} files")

    logger.info("\nFile Type Breakdown:")
    for file_type, info in stats["by_file_type"].items():
        logger.info(
            f"  {file_type}: {info['count']} files ({info['percentage']:.1f}%)")

    return processor


def test_metadata_extraction(processor):
    """Test metadata extraction from different file types."""
    logger.info("\n===== TESTING METADATA EXTRACTION =====")

    for grant_type, samples_list in processor.samples.items():
        logger.info(f"\nMetadata for {grant_type} samples:")

        for sample in samples_list:
            logger.info(
                f"\n  File: {sample['filename']} ({sample['file_type']})")

            if "metadata" in sample:
                for key, value in sample["metadata"].items():
                    if value is not None:
                        logger.info(f"    - {key}: {value}")
            else:
                logger.info("    No metadata available")


def test_style_prompt_generation(processor):
    """Test generating style-aware prompts."""
    logger.info("\n===== TESTING STYLE PROMPT GENERATION =====")

    # Test topics
    test_topics = [
        "AI-powered healthcare diagnostics",
        "Renewable energy in developing countries",
        "Educational technology for underprivileged communities",
        "Sustainable agriculture practices"
    ]

    # Test sections
    test_sections = [
        "Executive Summary",
        "Problem Statement",
        "Methodology",
        "Budget"
    ]

    # Generate prompts for different combinations
    for grant_type in GRANT_TYPES:
        if grant_type in processor.samples and processor.samples[grant_type]:
            topic = random.choice(test_topics)
            section = random.choice(test_sections)

            logger.info(
                f"\nGenerating prompt for {grant_type} grant, topic: {topic}, section: {section}")

            sample = processor.get_best_sample(grant_type, topic)
            if sample:
                logger.info(
                    f"Using sample: {sample['filename']} ({sample['file_type']})")

                start_time = time.time()
                prompt = processor.create_style_prompt(sample, topic, section)
                elapsed = time.time() - start_time

                logger.info(f"Generated prompt in {elapsed:.4f} seconds")
                logger.info(f"Prompt excerpt: {prompt[:150]}...")
            else:
                logger.info(f"No samples found for {grant_type}")


def test_batch_processing():
    """Test batch processing capability."""
    logger.info("\n===== TESTING BATCH PROCESSING =====")

    # Initialize processor with test directory
    processor = SampleGrantProcessor(samples_dir=str(TEST_DATA_DIR))

    # Time the batch processing
    start_time = time.time()
    processor.batch_process_samples(batch_size=3)
    elapsed = time.time() - start_time

    # Print summary
    total_samples = sum(len(s) for s in processor.samples.values())
    logger.info(
        f"Batch processed {total_samples} samples in {elapsed:.2f} seconds")

    # Print statistics
    stats = processor.get_sample_statistics()
    logger.info(f"Average word count: {stats['average_length']:.0f}")

    return processor


def test_error_handling():
    """Test error handling for unsupported files and edge cases."""
    logger.info("\n===== TESTING ERROR HANDLING =====")

    # Create a test directory for error cases
    error_dir = TEST_DATA_DIR / "error_tests"
    error_dir.mkdir(exist_ok=True)

    # Create an empty file
    empty_file = error_dir / "empty.txt"
    with open(empty_file, 'w') as f:
        pass

    # Create a file with unsupported extension
    unsupported_file = error_dir / "unsupported.xyz"
    with open(unsupported_file, 'w') as f:
        f.write("This is an unsupported file type")

    # Create a corrupt PDF if PDF support is available
    corrupt_pdf = None
    if PDF_SUPPORT:
        corrupt_pdf = error_dir / "corrupt.pdf"
        with open(corrupt_pdf, 'w') as f:
            f.write("This is not a valid PDF file")

    # Initialize processor with error test directory
    processor = SampleGrantProcessor(samples_dir=str(error_dir))

    # Test loading with error cases
    logger.info("Testing error handling during loading...")
    try:
        samples = processor.load_samples()
        logger.info(
            f"Loaded {sum(len(s) for s in samples.values())} samples (should be 0)")
    except Exception as e:
        logger.error(f"Unexpected error during loading: {e}")

    # Test extracting text from unsupported file
    logger.info("\nTesting unsupported file handling...")
    try:
        processor._extract_text_from_file(unsupported_file)
        logger.error("Failed: Unsupported file should have raised an error")
    except ValueError as e:
        logger.info(f"Success: Correctly handled unsupported file: {e}")
    except Exception as e:
        logger.error(f"Unexpected error type: {e}")

    # Test corrupt PDF if available
    if corrupt_pdf and PDF_SUPPORT:
        logger.info("\nTesting corrupt PDF handling...")
        try:
            processor._extract_text_from_file(corrupt_pdf)
            logger.error("Failed: Corrupt PDF should have raised an error")
        except Exception as e:
            logger.info(f"Success: Correctly handled corrupt PDF: {e}")

    # Clean up
    shutil.rmtree(error_dir)
    logger.info("Error handling tests completed")


def test_performance():
    """Test performance with different file types."""
    logger.info("\n===== TESTING PERFORMANCE =====")

    # Initialize processor with test directory
    processor = SampleGrantProcessor(samples_dir=str(TEST_DATA_DIR))

    # Test performance for each file type
    performance_stats = {}

    # Get all files by type
    txt_files = list(TEST_DATA_DIR.glob("**/*.txt"))
    pdf_files = list(TEST_DATA_DIR.glob("**/*.pdf")) if PDF_SUPPORT else []
    docx_files = list(TEST_DATA_DIR.glob("**/*.docx")) if DOCX_SUPPORT else []

    # Test text files
    if txt_files:
        start_time = time.time()
        for file_path in txt_files[:5]:  # Limit to 5 files
            try:
                content, metadata = processor._extract_text_from_file(
                    file_path)
                processor._validate_and_sanitize_text(content)
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
        txt_time = time.time() - start_time
        performance_stats["txt"] = txt_time / len(txt_files[:5])

    # Test PDF files
    if pdf_files:
        start_time = time.time()
        for file_path in pdf_files[:5]:  # Limit to 5 files
            try:
                content, metadata = processor._extract_text_from_file(
                    file_path)
                processor._validate_and_sanitize_text(content)
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
        pdf_time = time.time() - start_time
        performance_stats["pdf"] = pdf_time / len(pdf_files[:5])

    # Test Word files
    if docx_files:
        start_time = time.time()
        for file_path in docx_files[:5]:  # Limit to 5 files
            try:
                content, metadata = processor._extract_text_from_file(
                    file_path)
                processor._validate_and_sanitize_text(content)
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
        docx_time = time.time() - start_time
        performance_stats["docx"] = docx_time / len(docx_files[:5])

    # Print performance results
    logger.info("\nPerformance Results (average processing time per file):")
    for file_type, avg_time in performance_stats.items():
        logger.info(f"  {file_type}: {avg_time:.4f} seconds per file")

    # Compare performance
    if len(performance_stats) > 1:
        fastest = min(performance_stats.items(), key=lambda x: x[1])
        slowest = max(performance_stats.items(), key=lambda x: x[1])
        logger.info(f"\nFastest: {fastest[0]} ({fastest[1]:.4f}s)")
        logger.info(f"Slowest: {slowest[0]} ({slowest[1]:.4f}s)")
        logger.info(
            f"Slowest is {slowest[1]/fastest[1]:.1f}x slower than fastest")


def main():
    """Main test function."""
    logger.info("Starting enhanced sample processor tests...")

    # Create test files
    created_files = create_test_files()

    # Test file loading
    processor = test_file_loading()

    # Test metadata extraction
    test_metadata_extraction(processor)

    # Test style prompt generation
    test_style_prompt_generation(processor)

    # Test batch processing
    batch_processor = test_batch_processing()

    # Test error handling
    test_error_handling()

    # Test performance
    test_performance()

    logger.info("\n===== ALL TESTS COMPLETED =====")

    # Clean up test files (comment out to keep files for inspection)
    # shutil.rmtree(TEST_DATA_DIR)
    # logger.info(f"Cleaned up test directory: {TEST_DATA_DIR}")


if __name__ == "__main__":
    main()
