#!/usr/bin/env python


"""


Demo Enhanced Features





This interactive demonstration showcases the enhanced sample processor's


capabilities for handling multiple file types (PDF, Word, text) and


generating style-matched grant proposals.





Usage:


    python demo_enhanced_features.py [--create-samples] [--compare]





Features demonstrated:


1. Multi-format document processing (PDF, Word, text)


2. Metadata extraction and analysis


3. Style-matched prompt generation


4. Performance comparison between file types


5. Hybrid RAG integration


6. Batch processing capabilities


"""





import sys


import time


import argparse


import logging


import random


import shutil


from pathlib import Path


from typing import Dict, List





import importlib.util as _importlib_util
# Configure logging


logging.basicConfig(


    level=logging.INFO,


    format='%(asctime)s - %(levelname)s - %(message)s',


    datefmt='%H:%M:%S'


)


logger = logging.getLogger(__name__)





# Try importing the enhanced sample processor


try:


    from utils.sample_processor import SampleGrantProcessor, PDF_SUPPORT, DOCX_SUPPORT


except ImportError:


    logger.error(


        "Failed to import SampleGrantProcessor. Make sure you're in the project root directory.")


    sys.exit(1)





# Check for optional dependencies
PDF_AVAILABLE = _importlib_util.find_spec("pypdf") is not None





DOCX_AVAILABLE = _importlib_util.find_spec("docx") is not None





try:


    from reportlab.pdfgen import canvas


    from reportlab.lib.pagesizes import letter


    REPORTLAB_AVAILABLE = True


except ImportError:


    logger.warning(


        "reportlab not installed. Sample creation features will be limited.")


    REPORTLAB_AVAILABLE = False





# Constants


SAMPLE_DIR = Path("data/samples")


DEMO_DIR = Path("data/demo_samples")


GRANT_TYPES = ["research", "nonprofit",


               "technology", "healthcare", "education"]


SAMPLE_SECTIONS = [


    "EXECUTIVE SUMMARY",


    "PROBLEM STATEMENT",


    "METHODOLOGY",


    "EXPECTED OUTCOMES",


    "BUDGET",


    "TIMELINE"


]





# ASCII art banner


BANNER = """


â•”â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•—


â•‘                                                           â•‘


â•‘   Enhanced Sample Processor - Interactive Demo            â•‘


â•‘   â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•       â•‘


â•‘                                                           â•‘


â•‘   Multi-format Grant Sample Processing & Style Matching   â•‘


â•‘                                                           â•‘


â•šâ•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•


"""





# Helper functions for terminal output








def print_header(text):


    """Print a section header."""


    print(f"\n{'=' * 80}")


    print(f"  {text}")


    print(f"{'=' * 80}")








def print_subheader(text):


    """Print a subsection header."""


    print(f"\n{'-' * 60}")


    print(f"  {text}")


    print(f"{'-' * 60}")








def print_success(text):


    """Print a success message."""


    print(f"\nâœ“ {text}")








def print_warning(text):


    """Print a warning message."""


    print(f"\nâš  {text}")








def print_error(text):


    """Print an error message."""


    print(f"\nâœ— {text}")








def print_info(text):


    """Print an info message."""


    print(f"\nâ†’ {text}")








def print_table(headers, rows):


    """Print a simple ASCII table."""


    # Calculate column widths


    col_widths = [len(h) for h in headers]


    for row in rows:


        for i, cell in enumerate(row):


            col_widths[i] = max(col_widths[i], len(str(cell)))





    # Print headers


    header_row = " | ".join(


        h.ljust(col_widths[i]) for i, h in enumerate(headers))


    print(f"\n{header_row}")


    print("-" * len(header_row))





    # Print rows


    for row in rows:


        print(" | ".join(str(cell).ljust(


            col_widths[i]) for i, cell in enumerate(row)))








def wait_for_user():


    """Wait for user to press Enter to continue."""


    input("\nPress Enter to continue...")








def create_sample_files(demo_dir: Path = DEMO_DIR) -> Dict[str, List[Path]]:


    """Create sample files of different types for demonstration."""


    print_header("Creating Sample Files")





    # Create test directory structure


    if demo_dir.exists():


        print_info(f"Demo directory {demo_dir} already exists. Clearing it...")


        shutil.rmtree(demo_dir)





    # Create directories for each grant type


    created_files = {}


    for grant_type in GRANT_TYPES:


        type_dir = demo_dir / grant_type


        type_dir.mkdir(parents=True, exist_ok=True)


        created_files[grant_type] = []





        # Create text files


        print_info(f"Creating text files for {grant_type}...")


        for i in range(2):


            file_path = type_dir / f"sample_{grant_type}_{i+1}.txt"


            with open(file_path, 'w', encoding='utf-8') as f:


                f.write(f"# {grant_type.upper()} GRANT PROPOSAL {i+1}\n\n")





                # Add random sections


                for section in SAMPLE_SECTIONS:


                    f.write(f"\n{section}\n")


                    f.write(


                        f"This is sample content for the {section.lower()} section of a {grant_type} grant proposal.\n")


                    f.write(


                        "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.\n")


                    f.write(


                        "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.\n\n")





            created_files[grant_type].append(file_path)





        # Create PDF files if supported


        if REPORTLAB_AVAILABLE and PDF_AVAILABLE:


            print_info(f"Creating PDF files for {grant_type}...")


            try:


                for i in range(2):


                    file_path = type_dir / f"sample_{grant_type}_{i+1}.pdf"


                    c = canvas.Canvas(str(file_path), pagesize=letter)


                    c.setFont("Helvetica", 12)





                    # Add title


                    c.drawString(


                        72, 750, f"{grant_type.upper()} GRANT PROPOSAL {i+1}")





                    # Add sections


                    y_position = 700


                    for section in SAMPLE_SECTIONS:


                        c.drawString(72, y_position, section)


                        y_position -= 20


                        c.drawString(


                            72, y_position, f"This is sample content for the {section.lower()} section of a {grant_type} grant proposal.")


                        y_position -= 15


                        c.drawString(


                            72, y_position, "Lorem ipsum dolor sit amet, consectetur adipiscing elit.")


                        y_position -= 15


                        c.drawString(


                            72, y_position, "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris.")


                        y_position -= 30





                    c.save()


                    created_files[grant_type].append(file_path)





            except Exception as e:


                print_error(f"Error creating PDF files: {e}")


        else:


            print_warning(


                "PDF creation skipped: reportlab or pypdf not available")





        # Create Word documents if supported
        if DOCX_AVAILABLE:
            print_info(f"Creating Word documents for {grant_type}...")
            try:
                import docx
                for i in range(2):
                    file_path = type_dir / f"sample_{grant_type}_{i+1}.docx"
                    doc = docx.Document()
                    # Add title
                    doc.add_heading(f"{grant_type.upper()} GRANT PROPOSAL {i+1}", level=1)
                    # Add sections
                    for section in SAMPLE_SECTIONS:
                        doc.add_heading(section, level=2)
                        doc.add_paragraph(f"This is sample content for the {section.lower()} section of a {grant_type} grant proposal.")
                        p = doc.add_paragraph("Lorem ipsum dolor sit amet, consectetur adipiscing elit. ")
                        p.add_run("Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.").bold = True
                        doc.add_paragraph("Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.")
                    doc.save(file_path)
                    created_files[grant_type].append(file_path)
            except Exception as e:
                print_error(f"Error creating Word documents: {e}")
        else:
            print_warning("Word document creation skipped: python-docx not available")





    # Print summary


    total_files = sum(len(files) for files in created_files.values())


    print_success(f"Created {total_files} sample files in {demo_dir}")





    # Print breakdown


    file_types = {}


    for grant_type, files in created_files.items():


        for file in files:


            ext = file.suffix.lower()


            if ext not in file_types:


                file_types[ext] = 0


            file_types[ext] += 1





    print_info("File type breakdown:")


    for ext, count in file_types.items():


        print(f"  {ext}: {count} files")





    return created_files








def demo_file_loading(sample_dir: Path = DEMO_DIR):


    """Demonstrate loading files of different types."""


    print_header("Demonstrating File Loading")





    # Initialize processor with sample directory


    processor = SampleGrantProcessor(samples_dir=str(sample_dir))





    # Check what file types are supported


    print_info("Supported file extensions:")


    for ext in processor.supported_extensions:


        print(f"  {ext}")





    # Time the loading process


    print_info("Loading samples...")


    start_time = time.time()


    samples = processor.load_samples(show_progress=True)


    elapsed = time.time() - start_time





    # Print summary


    total_samples = sum(len(s) for s in samples.values())


    print_success(f"Loaded {total_samples} samples in {elapsed:.2f} seconds")





    # Print breakdown by file type


    stats = processor.get_sample_statistics()


    print_subheader("Sample Statistics")





    # Print grant type statistics


    headers = ["Grant Type", "Count", "Avg Words", "File Types"]


    rows = []


    for grant_type, info in stats["by_type"].items():


        file_types_str = ", ".join(


            f"{ft} ({count})" for ft, count in info["file_types"].items())


        rows.append([


            grant_type,


            info["count"],


            f"{info['avg_words']:.0f}",


            file_types_str


        ])





    print_table(headers, rows)





    # Print file type statistics


    print_subheader("File Type Statistics")


    headers = ["File Type", "Count", "Percentage"]


    rows = []


    for file_type, info in stats["by_file_type"].items():


        rows.append([


            file_type,


            info["count"],


            f"{info['percentage']:.1f}%"


        ])





    print_table(headers, rows)





    return processor








def demo_metadata_extraction(processor):


    """Demonstrate metadata extraction from different file types."""


    print_header("Demonstrating Metadata Extraction")





    # Select a few samples to show metadata


    samples_to_show = []


    for grant_type, samples_list in processor.samples.items():


        # Get one sample of each file type if available


        file_types_shown = set()


        for sample in samples_list:


            file_type = sample.get("file_type", "unknown")


            if file_type not in file_types_shown:


                samples_to_show.append(sample)


                file_types_shown.add(file_type)


                if len(file_types_shown) >= 3:  # Show at most 3 samples per grant type


                    break





    # Show metadata for selected samples


    for sample in samples_to_show:


        print_subheader(


            f"Metadata for {sample['filename']} ({sample.get('file_type', 'unknown')})")





        if "metadata" in sample:


            # Format metadata for display


            for key, value in sample["metadata"].items():


                if value is not None:


                    print(f"  {key}: {value}")


        else:


            print_warning("No metadata available")





    print_info("Metadata extraction capabilities depend on file type:")


    print("  - PDF: Author, creation date, page count, producer")


    print("  - DOCX: Author, created/modified dates, title, paragraph count")


    print("  - TXT: Basic file system metadata (created/modified dates, size)")








def demo_style_prompt_generation(processor):


    """Demonstrate generating style-aware prompts."""


    print_header("Demonstrating Style-Aware Prompt Generation")





    # Test topics and sections


    test_topics = [


        "AI-powered healthcare diagnostics",


        "Renewable energy in developing countries",


        "Educational technology for underprivileged communities"


    ]





    test_sections = [


        "Executive Summary",


        "Problem Statement",


        "Methodology"


    ]





    # Let user choose a topic and section


    print_info("Available topics:")


    for i, topic in enumerate(test_topics):


        print(f"  {i+1}. {topic}")





    topic_idx = input("\nSelect a topic (1-3) or enter your own: ")


    try:


        topic_idx = int(topic_idx) - 1


        if 0 <= topic_idx < len(test_topics):


            topic = test_topics[topic_idx]


        else:


            topic = topic_idx


    except ValueError:


        topic = topic_idx





    print_info("Available sections:")


    for i, section in enumerate(test_sections):


        print(f"  {i+1}. {section}")





    section_idx = input("\nSelect a section (1-3) or enter your own: ")


    try:


        section_idx = int(section_idx) - 1


        if 0 <= section_idx < len(test_sections):


            section = test_sections[section_idx]


        else:


            section = section_idx


    except ValueError:


        section = section_idx





    # Let user choose a grant type


    available_grant_types = [


        gt for gt in processor.samples.keys() if processor.samples[gt]]





    if not available_grant_types:


        print_error("No grant samples available. Please create samples first.")


        return





    print_info("Available grant types:")


    for i, grant_type in enumerate(available_grant_types):


        print(f"  {i+1}. {grant_type}")





    grant_idx = input(


        "\nSelect a grant type (1-{0}): ".format(len(available_grant_types)))


    try:


        grant_idx = int(grant_idx) - 1


        if 0 <= grant_idx < len(available_grant_types):


            grant_type = available_grant_types[grant_idx]


        else:


            grant_type = available_grant_types[0]


    except ValueError:


        grant_type = available_grant_types[0]





    print_info(


        f"Generating prompt for {grant_type} grant, topic: {topic}, section: {section}")





    # Get sample and generate prompt


    sample = processor.get_best_sample(grant_type, topic)


    if sample:


        print_success(


            f"Using sample: {sample['filename']} ({sample.get('file_type', 'unknown')})")





        start_time = time.time()


        prompt = processor.create_style_prompt(sample, topic, section)


        elapsed = time.time() - start_time





        print_info(f"Generated prompt in {elapsed:.4f} seconds")





        # Show the prompt


        print_subheader("Generated Style-Aware Prompt")


        print(prompt)





        # Explain what happened


        print_info("The prompt above combines:")


        print("  1. The structure and style from the sample grant")


        print("  2. Your specific topic and section requirements")


        print("  3. Instructions to maintain the same level of detail and professional language")





    else:


        print_error(f"No samples found for {grant_type}")








def demo_performance_comparison(sample_dir: Path = DEMO_DIR):


    """Demonstrate performance comparison between file types."""


    print_header("Performance Comparison Between File Types")





    # Initialize processor with sample directory


    processor = SampleGrantProcessor(samples_dir=str(sample_dir))





    # Get all files by type


    txt_files = list(sample_dir.glob("**/*.txt"))


    pdf_files = list(sample_dir.glob("**/*.pdf")) if PDF_SUPPORT else []


    docx_files = list(sample_dir.glob("**/*.docx")) if DOCX_SUPPORT else []





    performance_stats = {}


    file_counts = {


        "txt": len(txt_files),


        "pdf": len(pdf_files),


        "docx": len(docx_files)


    }





    print_info("Testing performance for each file type...")





    # Test text files


    if txt_files:


        print_info(f"Processing {min(5, len(txt_files))} text files...")


        start_time = time.time()


        for file_path in txt_files[:5]:  # Limit to 5 files


            try:


                content, metadata = processor._extract_text_from_file(


                    file_path)


                processor._validate_and_sanitize_text(content)


            except Exception as e:


                print_error(f"Error processing {file_path}: {e}")


        txt_time = time.time() - start_time


        performance_stats["txt"] = txt_time / min(5, len(txt_files))





    # Test PDF files


    if pdf_files:


        print_info(f"Processing {min(5, len(pdf_files))} PDF files...")


        start_time = time.time()


        for file_path in pdf_files[:5]:  # Limit to 5 files


            try:


                content, metadata = processor._extract_text_from_file(


                    file_path)


                processor._validate_and_sanitize_text(content)


            except Exception as e:


                print_error(f"Error processing {file_path}: {e}")


        pdf_time = time.time() - start_time


        performance_stats["pdf"] = pdf_time / min(5, len(pdf_files))





    # Test Word files


    if docx_files:


        print_info(f"Processing {min(5, len(docx_files))} Word files...")


        start_time = time.time()


        for file_path in docx_files[:5]:  # Limit to 5 files


            try:


                content, metadata = processor._extract_text_from_file(


                    file_path)


                processor._validate_and_sanitize_text(content)


            except Exception as e:


                print_error(f"Error processing {file_path}: {e}")


        docx_time = time.time() - start_time


        performance_stats["docx"] = docx_time / min(5, len(docx_files))





    # Print performance results


    print_subheader("Performance Results")





    headers = ["File Type", "Avg. Time (s)", "Files Found", "Relative Speed"]


    rows = []





    # Find the fastest for relative comparison


    if performance_stats:


        fastest = min(performance_stats.values())





        for file_type, avg_time in performance_stats.items():


            relative = avg_time / fastest if fastest > 0 else 0


            rows.append([


                file_type,


                f"{avg_time:.4f}",


                file_counts[file_type],


                f"{relative:.1f}x" if relative > 1 else "1.0x (fastest)"


            ])





        print_table(headers, rows)





        # Print summary


        fastest_type = min(performance_stats.items(), key=lambda x: x[1])[0]


        slowest_type = max(performance_stats.items(), key=lambda x: x[1])[0]


        print_info(f"Fastest: {fastest_type} files")


        print_info(


            f"Slowest: {slowest_type} files ({performance_stats[slowest_type]/performance_stats[fastest_type]:.1f}x slower)")





        # Print throughput estimate


        print_subheader("Estimated Throughput")


        for file_type, avg_time in performance_stats.items():


            throughput = 3600 / avg_time  # files per hour


            print(f"  {file_type}: {throughput:.0f} files/hour")


    else:


        print_warning(


            "No performance data available. Please create sample files first.")








def demo_batch_processing(sample_dir: Path = DEMO_DIR):


    """Demonstrate batch processing capability."""


    print_header("Demonstrating Batch Processing")





    # Initialize processor with sample directory


    processor = SampleGrantProcessor(samples_dir=str(sample_dir))





    # Get total file count


    total_files = 0


    for ext in processor.supported_extensions:


        total_files += len(list(sample_dir.glob(f"**/*{ext}")))





    if total_files == 0:


        print_warning("No sample files found. Please create samples first.")


        return





    # Ask for batch size


    batch_size = input(f"\nEnter batch size (1-{total_files}, default: 3): ")


    try:


        batch_size = int(batch_size)


        if batch_size < 1 or batch_size > total_files:


            batch_size = 3


    except ValueError:


        batch_size = 3





    print_info(f"Processing {total_files} files in batches of {batch_size}...")





    # Time the batch processing


    start_time = time.time()


    processor.batch_process_samples(batch_size=batch_size)


    elapsed = time.time() - start_time





    # Print summary


    total_samples = sum(len(s) for s in processor.samples.values())


    print_success(


        f"Batch processed {total_samples} samples in {elapsed:.2f} seconds")





    # Calculate and print throughput


    throughput = total_samples / elapsed


    print_info(f"Processing throughput: {throughput:.2f} files/second")


    print_info(f"Estimated time for 1000 files: {1000/throughput:.1f} seconds")





    # Print statistics


    stats = processor.get_sample_statistics()


    print_subheader("Sample Statistics After Batch Processing")


    print(f"Total samples: {stats['total_samples']}")


    print(f"Average word count: {stats['average_length']:.0f}")





    # Show file type breakdown


    print_info("File type breakdown:")


    for file_type, info in stats["by_file_type"].items():


        print(


            f"  {file_type}: {info['count']} files ({info['percentage']:.1f}%)")





    return processor








def demo_hybrid_integration(processor):


    """Demonstrate hybrid integration with RAG."""


    print_header("Demonstrating Hybrid Integration")





    # Check if we can import the necessary modules


    try:


        from utils.rag_utils import query_context


        rag_available = True


    except ImportError:


        print_warning(


            "RAG utilities not available. Showing simulated integration.")


        rag_available = False





    # Create a simple hybrid integration demo


    print_info("The hybrid approach combines:")


    print("  1. Style-matched prompts from sample grants")


    print("  2. Relevant context from the RAG vector database")


    print("  3. User-specific requirements")





    # Let user input a topic


    topic = input(


        "\nEnter a grant topic (e.g., 'AI-powered healthcare diagnostics'): ")


    if not topic:


        topic = "AI-powered healthcare diagnostics"





    # Select a grant type and section


    available_grant_types = [


        gt for gt in processor.samples.keys() if processor.samples[gt]]


    if not available_grant_types:


        print_error("No grant samples available. Please create samples first.")


        return





    grant_type = random.choice(available_grant_types)


    section = "Problem Statement"





    print_info(


        f"Creating hybrid prompt for {grant_type} grant about '{topic}'")





    # Get sample-based prompt


    sample = processor.get_best_sample(grant_type, topic)


    if sample:


        print_success(


            f"Using sample: {sample['filename']} ({sample.get('file_type', 'unknown')})")


        sample_prompt = processor.create_style_prompt(sample, topic, section)


        print_info("Generated sample-based prompt")


    else:


        print_warning(f"No samples found for {grant_type}")


        sample_prompt = f"Write a {section} for a {grant_type} grant proposal about '{topic}'."





    # Simulate or get RAG context


    if rag_available:


        try:


            # Create a query based on the topic and section


            query = f"{section} for {grant_type} grant about {topic}"





            # Get relevant context from RAG


            print_info("Retrieving relevant context from RAG...")


            rag_results = query_context(query, k=3)





            if rag_results:


                # Extract and format the context


                rag_context = ""


                for i, doc in enumerate(rag_results):


                    rag_context += f"\nRELEVANT CONTEXT {i+1}:\n{doc.page_content[:200]}...\n"





                print_success(


                    f"Retrieved {len(rag_results)} relevant documents from RAG")


            else:


                print_warning("No relevant context found in RAG")


                rag_context = "No relevant context available."


        except Exception as e:


            print_error(f"Error retrieving RAG context: {e}")


            rag_context = "Error retrieving context from RAG system."


    else:


        # Simulate RAG context


        print_info("Simulating RAG context retrieval...")


        rag_context = f"""


        RELEVANT CONTEXT 1:


        Previous successful {grant_type} grants have emphasized the importance of clear problem statements


        that highlight the gap being addressed and the potential impact of the solution.


        


        RELEVANT CONTEXT 2:


        When writing about {topic}, it's important to cite recent statistics and research findings


        to establish credibility and urgency.


        """


        print_success("Retrieved simulated RAG context")





    # Combine sample prompt with RAG context


    hybrid_prompt = f"""


    {sample_prompt}


    


    Additionally, consider the following relevant information:


    {rag_context}


    


    Remember to maintain a professional tone and focus specifically on {topic}.


    """





    # Show the hybrid prompt


    print_subheader("Generated Hybrid Prompt")


    print(hybrid_prompt)





    # Explain the benefits


    print_subheader("Benefits of Hybrid Approach")


    print("1. Style consistency from sample grants")


    print("2. Domain knowledge from RAG system")


    print("3. Personalization based on user requirements")


    print("4. Better factual accuracy and relevance")


    print("5. Improved overall quality of generated content")








def demo_comparison_with_original():


    """Demonstrate comparison with the original processor."""


    print_header("Comparing Enhanced vs Original Processor")





    # Create a table showing the differences


    print_subheader("Feature Comparison")





    headers = ["Feature", "Original", "Enhanced"]


    rows = [


        ["File types supported", "Text only", "Text, PDF, Word"],


        ["Metadata extraction", "Basic", "Comprehensive"],


        ["Section detection", "Basic patterns", "Advanced patterns"],


        ["Processing mode", "All at once", "Batch or all at once"],


        ["Progress tracking", "No", "Yes"],


        ["Performance metrics", "No", "Yes"],


        ["File validation", "Basic", "Advanced sanitization"],


        ["Statistics", "Basic counts", "Detailed by file type"],


        ["RAG integration", "Manual", "Automated"],


        ["Error handling", "Basic", "Comprehensive"],


    ]





    print_table(headers, rows)





    # Show a visual representation of capabilities


    print_subheader("Capability Expansion")





    print("Original Processor:")


    print("  [â–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆ] Text files")


    print("  [                    ] PDF files")


    print("  [                    ] Word documents")


    print("  [â–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆ              ] Metadata extraction")


    print("  [â–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆ            ] Section detection")


    print("  [                    ] Batch processing")


    print("  [â–ˆâ–ˆ                  ] Error handling")





    print("\nEnhanced Processor:")


    print("  [â–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆ] Text files")


    print("  [â–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆ] PDF files")


    print("  [â–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆ] Word documents")


    print("  [â–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆ    ] Metadata extraction")


    print("  [â–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆ  ] Section detection")


    print("  [â–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆ] Batch processing")


    print("  [â–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆâ–ˆ  ] Error handling")





    # Explain the benefits


    print_subheader("Key Benefits")


    print("1. Multi-format support expands usable grant samples by 300%+")


    print("2. Metadata extraction enables better sample organization and filtering")


    print("3. Batch processing enables handling large document collections efficiently")


    print("4. Improved error handling ensures robustness with varied input quality")


    print("5. Detailed statistics provide insights into your grant sample collection")


    print("6. Seamless integration with RAG improves context retrieval")








def main():


    """Main function to run the demonstration."""


    # Parse command line arguments


    parser = argparse.ArgumentParser(


        description="Demo Enhanced Sample Processor Features")


    parser.add_argument("--create-samples", action="store_true",


                        help="Create sample files for demonstration")


    parser.add_argument("--compare", action="store_true",


                        help="Show comparison with original processor")


    parser.add_argument("--sample-dir", default=str(DEMO_DIR),


                        help="Directory for sample files")


    args = parser.parse_args()





    # Print banner


    print(BANNER)





    # Show system info


    print_info("System Information:")


    print(f"  Python version: {sys.version.split()[0]}")


    print(


        f"  PDF support: {'Available' if PDF_AVAILABLE else 'Not available'}")


    print(


        f"  Word support: {'Available' if DOCX_AVAILABLE else 'Not available'}")


    print(


        f"  Sample creation: {'Available' if REPORTLAB_AVAILABLE else 'Limited'}")





    # Create sample files if requested


    sample_dir = Path(args.sample_dir)


    if args.create_samples:


        create_sample_files(sample_dir)


        wait_for_user()





    # Show comparison if requested


    if args.compare:


        demo_comparison_with_original()


        wait_for_user()





    # Main demo flow


    while True:


        print_header("Demo Menu")


        print("1. Create sample files (PDF, Word, Text)")


        print("2. Load and analyze samples")


        print("3. Extract and view metadata")


        print("4. Generate style-matched prompts")


        print("5. Compare performance between file types")


        print("6. Demonstrate batch processing")


        print("7. Show hybrid integration with RAG")


        print("8. Compare with original processor")


        print("0. Exit")





        choice = input("\nEnter your choice (0-8): ")





        if choice == "0":


            print_info("Exiting demo. Thank you!")


            break


        elif choice == "1":


            create_sample_files(sample_dir)


        elif choice == "2":


            processor = demo_file_loading(sample_dir)


        elif choice == "3":


            try:


                processor


            except NameError:


                processor = demo_file_loading(sample_dir)


            demo_metadata_extraction(processor)


        elif choice == "4":


            try:


                processor


            except NameError:


                processor = demo_file_loading(sample_dir)


            demo_style_prompt_generation(processor)


        elif choice == "5":


            demo_performance_comparison(sample_dir)


        elif choice == "6":


            processor = demo_batch_processing(sample_dir)


        elif choice == "7":


            try:


                processor


            except NameError:


                processor = demo_file_loading(sample_dir)


            demo_hybrid_integration(processor)


        elif choice == "8":


            demo_comparison_with_original()


        else:


            print_warning("Invalid choice. Please try again.")





        wait_for_user()








if __name__ == "__main__":


    try:


        main()


    except KeyboardInterrupt:


        print_info("\nDemo interrupted. Exiting gracefully.")


    except Exception as e:


        print_error(f"An error occurred: {e}")


        import traceback


        traceback.print_exc()


